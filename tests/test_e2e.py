"""End-to-end smoke test: generate synthetic `.docx` fixtures, run them
through LibreOffice + pypdf, verify the final binder layout.

This is the *only* place we exercise the LibreOffice + trim-heuristic +
layout pipeline together. It's slower than the unit tests (LibreOffice
startup is ~5s per file), so the CI workflow runs it in a separate job.

Skips automatically when `soffice` isn't on PATH and the macOS app bundle
isn't present, so the test is harmless to run locally on a dev machine
without LibreOffice.
"""

from __future__ import annotations

import os
import shutil
from pathlib import Path

import pytest
from build_binder import (
    Config,
    convert_to_pdf,
    merge_into_binder,
    plan_layout,
    resolve_soffice,
)
from docx import Document
from pypdf import PdfReader

pytestmark = pytest.mark.e2e


# ---------------------------------------------------------------------------
# soffice discovery — skip the entire module if LibreOffice isn't available.
# ---------------------------------------------------------------------------


def _find_soffice() -> Path | None:
    candidates = [
        Path("/Applications/LibreOffice.app/Contents/MacOS/soffice"),
    ]
    which = shutil.which("soffice")
    if which:
        candidates.insert(0, Path(which))
    for c in candidates:
        if c.exists() and os.access(c, os.X_OK):
            return c
    return None


SOFFICE = _find_soffice()

if SOFFICE is None:
    pytest.skip("LibreOffice (`soffice`) not installed", allow_module_level=True)


# ---------------------------------------------------------------------------
# Fixture generation
# ---------------------------------------------------------------------------


# Fixtures use explicit page breaks rather than relying on content overflow
# so the resulting PDFs are deterministic regardless of LibreOffice's font
# substitution / margin defaults on the current host.


def _make_single_page_docx(path: Path, title: str) -> None:
    doc = Document()
    section = doc.sections[0]
    section.header.paragraphs[0].text = f"Test Header — {title}"
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Verse for {title}. " * 8)
    doc.save(path)


def _make_two_page_docx(path: Path, title: str) -> None:
    """Two distinct pages with an explicit page break between them."""
    doc = Document()
    section = doc.sections[0]
    section.header.paragraphs[0].text = f"Test Header — {title}"
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Verse 1 for {title}. " + ("Page one lyrics. " * 8))
    doc.add_page_break()
    doc.add_paragraph(f"Verse 2 for {title}. " + ("Page two lyrics distinct words. " * 8))
    doc.save(path)


def _make_trailing_chrome_docx(path: Path, title: str) -> None:
    """Page 1 has substantive content. Page 2 (after an explicit page break)
    has only the header line text repeated as a body paragraph — every word
    is already present on page 1 via the repeated header, so the trim
    heuristic should drop page 2.
    """
    chrome_line = f"Test Header — {title}"
    doc = Document()
    section = doc.sections[0]
    section.header.paragraphs[0].text = chrome_line
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Verse for {title}. " + ("Lyrics line content. " * 10))
    doc.add_page_break()
    # Page-2 body intentionally duplicates the header text. After rendering,
    # page 2 will contain only the (repeated) header + this paragraph, which
    # also matches the header verbatim — zero unique words vs page 1.
    doc.add_paragraph(chrome_line)
    doc.save(path)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _build_config(tmp_path: Path) -> Config:
    out_dir = tmp_path / "out"
    out_dir.mkdir()
    return Config(
        chord_sheets_dir=tmp_path,  # not actually used in this test
        output_dir=out_dir,
        soffice_path=SOFFICE,  # type: ignore[arg-type]
        fuzzy_threshold=0.75,
    )


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------


def test_soffice_can_convert_a_docx(tmp_path: Path):
    """Sanity check: LibreOffice headless conversion works in this env."""
    src = tmp_path / "Single.docx"
    _make_single_page_docx(src, "Single Page Song")
    pdf = convert_to_pdf(src, tmp_path, SOFFICE)  # type: ignore[arg-type]
    assert pdf.exists()
    assert pdf.suffix == ".pdf"
    reader = PdfReader(str(pdf))
    assert len(reader.pages) >= 1


def test_layout_one_two_one_produces_expected_spread(tmp_path: Path):
    """End-to-end against the canonical user example: [2-page, 1-page, 1-page]
    must come out as page 1 blank, pages 2-3 the two-pager, page 4 a 1-pager,
    page 5 another 1-pager. The blank slot at the start is the key invariant.
    """
    sources = []
    a = tmp_path / "Two.docx"
    _make_two_page_docx(a, "Song A")
    sources.append(a)
    b = tmp_path / "One B.docx"
    _make_single_page_docx(b, "Song B")
    sources.append(b)
    c = tmp_path / "One C.docx"
    _make_single_page_docx(c, "Song C")
    sources.append(c)

    pdf_dir = tmp_path / "pdfs"
    pdf_dir.mkdir()
    pdfs = []
    for src in sources:
        pdfs.append((src, convert_to_pdf(src, pdf_dir, SOFFICE)))  # type: ignore[arg-type]

    out = tmp_path / "binder.pdf"
    placements = merge_into_binder(pdfs, out)

    assert len(placements) == 3
    # Song A: 2 pages, blank inserted before it, lands on 2-3.
    assert placements[0].pages == 2
    assert placements[0].binder_start == 2
    assert placements[0].binder_end == 3
    # Song B: 1 page, lands on page 4.
    assert placements[1].pages == 1
    assert placements[1].binder_start == 4
    # Song C: 1 page, lands on page 5.
    assert placements[2].pages == 1
    assert placements[2].binder_start == 5

    # Final binder is exactly 5 pages.
    reader = PdfReader(str(out))
    assert len(reader.pages) == 5


def test_trailing_chrome_page_is_trimmed_end_to_end(tmp_path: Path):
    """A .docx with a real page 1 and a page-break-only page 2 should produce
    a 2-page PDF after LibreOffice conversion, then get auto-trimmed back to
    1 effective page by the merge step. This is the regression guard against
    bugs in the chrome-trim heuristic against real LibreOffice output.
    """
    src = tmp_path / "Trailing.docx"
    _make_trailing_chrome_docx(src, "Trailing Chrome Song")

    pdf_dir = tmp_path / "pdfs"
    pdf_dir.mkdir()
    pdf = convert_to_pdf(src, pdf_dir, SOFFICE)  # type: ignore[arg-type]

    raw_reader = PdfReader(str(pdf))
    # LibreOffice should render the page break as a separate trailing page
    # containing only the repeated header.
    assert len(raw_reader.pages) == 2, (
        f"expected 2 pages from LibreOffice, got {len(raw_reader.pages)}"
    )

    out = tmp_path / "binder.pdf"
    placements = merge_into_binder([(src, pdf)], out)
    assert len(placements) == 1
    assert placements[0].pages == 1, (
        f"trim should leave 1 effective page, got {placements[0].pages}"
    )
    assert placements[0].trimmed == 1
    assert placements[0].raw_pages == 2
    # Single-page song lands on page 1 (no blank).
    assert placements[0].binder_start == 1
    assert placements[0].binder_end == 1

    reader = PdfReader(str(out))
    assert len(reader.pages) == 1


def test_plan_layout_matches_end_to_end_placement(tmp_path: Path):
    """The pure plan_layout() must agree with what merge_into_binder() actually
    produces. Property regression test against drift between the two.
    """
    docs = []
    a = tmp_path / "A.docx"
    _make_two_page_docx(a, "A")
    docs.append(a)
    b = tmp_path / "B.docx"
    _make_single_page_docx(b, "B")
    docs.append(b)
    c = tmp_path / "C.docx"
    _make_two_page_docx(c, "C")
    docs.append(c)

    pdf_dir = tmp_path / "pdfs"
    pdf_dir.mkdir()
    pdfs = [(d, convert_to_pdf(d, pdf_dir, SOFFICE)) for d in docs]  # type: ignore[arg-type]
    out = tmp_path / "binder.pdf"
    placements = merge_into_binder(pdfs, out)

    pure_layout = plan_layout([p.pages for p in placements])
    assert [(e.binder_start, e.binder_end) for e in pure_layout] == [
        (p.binder_start, p.binder_end) for p in placements
    ]


def test_resolve_soffice_finds_installed_soffice():
    """Sanity check: resolve_soffice() with no configured path returns
    something executable on this machine."""
    found = resolve_soffice("")
    assert found.exists()
    assert os.access(found, os.X_OK)
